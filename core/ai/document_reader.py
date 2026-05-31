"""Best-effort text extraction for Document file_key blobs stored on R2/S3."""

from __future__ import annotations

import io
import logging
from typing import Any

from django.core.files.storage import default_storage

logger = logging.getLogger(__name__)


_PDF_TYPES = {"application/pdf"}
_DOCX_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
_PLAIN_TYPES = {
    "text/plain",
    "text/markdown",
    "text/csv",
    "application/json",
    "application/xml",
    "text/xml",
    "text/html",
}


def _ext_from_name(name: str | None) -> str:
    if not name:
        return ""
    name = name.lower()
    for ext in (".pdf", ".docx", ".txt", ".md", ".csv", ".json", ".xml", ".html"):
        if name.endswith(ext):
            return ext
    return ""


def _kind(mime_type: str | None, file_name: str | None) -> str:
    mt = (mime_type or "").lower().split(";", 1)[0].strip()
    if mt in _PDF_TYPES:
        return "pdf"
    if mt in _DOCX_TYPES:
        return "docx"
    if mt in _PLAIN_TYPES or mt.startswith("text/"):
        return "plain"
    ext = _ext_from_name(file_name)
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext in (".txt", ".md", ".csv", ".json", ".xml", ".html"):
        return "plain"
    return "unknown"


def _read_pdf(blob: bytes, max_pages: int) -> tuple[str, dict[str, Any]]:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise RuntimeError("pypdf not installed; run pip install pypdf") from exc
    reader = PdfReader(io.BytesIO(blob))
    total_pages = len(reader.pages)
    pages_read = min(total_pages, max_pages)
    chunks: list[str] = []
    for i in range(pages_read):
        try:
            chunks.append(reader.pages[i].extract_text() or "")
        except Exception as exc:  # noqa: BLE001 — per-page extract failure is non-fatal
            logger.warning("[AI] pdf.page_extract_failed page=%s err=%s", i, exc)
            chunks.append("")
    text = "\n\n".join(c.strip() for c in chunks if c and c.strip())
    return text, {"pages_total": total_pages, "pages_read": pages_read}


def _read_docx(blob: bytes) -> tuple[str, dict[str, Any]]:
    try:
        from docx import Document as DocxDocument
    except Exception as exc:
        raise RuntimeError("python-docx not installed") from exc
    doc = DocxDocument(io.BytesIO(blob))
    paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # Pull table cells too — leave/HR docs often use tables.
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
            if row_text.strip():
                paras.append(row_text)
    text = "\n".join(paras)
    return text, {"paragraphs": len(paras)}


def _read_plain(blob: bytes) -> tuple[str, dict[str, Any]]:
    for encoding in ("utf-8", "latin-1"):
        try:
            return blob.decode(encoding), {"encoding": encoding}
        except UnicodeDecodeError:
            continue
    return blob.decode("utf-8", errors="replace"), {"encoding": "utf-8-replace"}


def extract_text(
    *,
    file_key: str,
    mime_type: str | None,
    file_name: str | None,
    max_pages: int = 20,
    max_chars: int = 12000,
) -> dict[str, Any]:
    """Fetch a document from storage and return extracted plain text.

    Returns dict with keys:
      kind, text, truncated (bool), char_count, plus kind-specific metadata.
    Raises RuntimeError on unsupported types or unreadable blobs.
    """
    if not file_key:
        raise RuntimeError("Document has no file_key.")
    kind = _kind(mime_type, file_name)
    if kind == "unknown":
        raise RuntimeError(
            f"Unsupported document type (mime={mime_type!r}, name={file_name!r}). "
            "Supported: PDF, DOCX, plain text."
        )
    try:
        fh = default_storage.open(file_key, "rb")
        try:
            blob = fh.read()
        finally:
            fh.close()
    except Exception as exc:
        raise RuntimeError(f"Failed to fetch document blob: {exc}") from exc

    if kind == "pdf":
        text, meta = _read_pdf(blob, max_pages=max_pages)
    elif kind == "docx":
        text, meta = _read_docx(blob)
    else:
        text, meta = _read_plain(blob)

    text = (text or "").strip()
    truncated = len(text) > max_chars
    if truncated:
        text = text[:max_chars] + "\n…[truncated]"
    return {
        "kind": kind,
        "text": text,
        "truncated": truncated,
        "char_count": len(text),
        "byte_size": len(blob),
        **meta,
    }
